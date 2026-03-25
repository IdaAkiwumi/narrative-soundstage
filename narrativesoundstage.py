"""
PROJECT: Narrative Soundstage
VERSION: 1.0.0
AUTHOR: Ida Akiwumi
ROLE: Product Architect | Narrative Strategist | Screenwriter
TECH STACK: Python, Streamlit, Edge-TTS, Asyncio, Regex

DESCRIPTION:
An intelligent script-to-performance engine designed for screenwriters and producers.
This tool automates voice-over table reads, maintains character consistency through
dynamic casting, and provides a real-time "active-line" prompter for editing.

IDEAL FOR:
- Film & Television Pre-production, Script Development, Screenplay Revisions & Table Reads
- Game Design Narrative Logic
- Legal Thriller & Action Comedy Script Development
- Data Storytelling & AI-Assisted Workflows
"""

__author__ = "Ida Akiwumi"
__version__ = "1.0.0"
__license__ = "Proprietary"
__status__ = "Production / Portfolio"

import streamlit as st
from docx import Document
from docx.shared import Pt
import re
import asyncio
import edge_tts
import io
import base64
import time
import streamlit.components.v1 as components
import random

# --- UI SETUP: MUST COME BEFORE OTHER STREAMLIT COMMANDS ---
st.set_page_config(
    page_title="Narrative Soundstage | Designed by Ida Akiwumi",
    page_icon="🎭",
    layout="wide"
)

# Helps the selected/highlighted line stay visible in the textarea
HIGHLIGHT_SETTLE_TIME = 0.85
MIN_LINE_DISPLAY_TIME = 1.5

# Faster transition timing for non-dialogue speech cues
CAST_NAME_WAIT = 0.45
PARENTHETICAL_WAIT = 0.60

# Add top padding to uploaded scripts only so the first highlighted lines
# are less likely to hide at the top of the white script viewer.
UPLOAD_TOP_PADDING_LINES = 10

PLACEHOLDER_SCRIPT = f"""TITLE: NARRATIVE SOUNDSTAGE
AUTHOR: {__author__}
VERSION: {__version__}

INT. STUDIO - DAY

The screen is dark. A yellow glow emanates from the corner.

IDA (V.O.)
Welcome to the Soundstage. This tool is an intelligent script-to-performance engine.

NARRATOR
Designed for screenwriters and producers, it automates table reads and maintains character consistency.

PRODUCER
It handles everything from legal thrillers to action comedies.

IDA
To begin, simply upload your .docx script using the button above.
"""

FREE_VOICES = {
    # --- MALE ---
    "Guy (Male - US)": "en-US-GuyNeural",
    "Christopher (Male - Deep)": "en-US-ChristopherNeural",
    "Eric (Male - UK Narrative)": "en-GB-RyanNeural",
    "Thomas (Male - UK Formal)": "en-GB-ThomasNeural",
    "Natascha (Male - AU Aboriginal Hint)": "en-AU-WilliamNeural",
    "Abeo (Male - African/NG)": "en-NG-AbeoNeural",
    "Hamdan (Male - Arabic/UAE)": "ar-AE-HamdanNeural",
    "Asad (Male - Asian/PK)": "ur-PK-AsadNeural",
    "Liam (Male - Canadian)": "en-CA-LiamNeural",
    "Andrew (Male - Energetic/Afro-Latino)": "en-US-AndrewNeural",
    "Brian (Male - Casual US/Afro-Latino)": "en-US-BrianNeural",
    "Luke (Male - Caribbean/SA Hint)": "en-ZA-LukeNeural",
    "Chilemba (Male - East African)": "en-KE-ChilembaNeural",
    "Themba (Male - Zulu/Southern African)": "zu-ZA-ThembaNeural",
    "Gonzalo (Male - Latino-Strong)": "es-CO-GonzaloNeural",
    "Emilio (Male - Latino)": "es-DO-EmilioNeural",
    "Dmitry (Male - Russian)": "ru-RU-DmitryNeural",
    "Henri (Male - French)": "fr-FR-HenriNeural",
    "Conrad (Male - German)": "de-DE-ConradNeural",
    "Connor (Male - Irish)": "en-IE-ConnorNeural",

    # --- FEMALE ---
    "Aria (Female - US Pro)": "en-US-AriaNeural",
    "Jenny (Female - US Friendly)": "en-US-JennyNeural",
    "Sonia (Female - UK Soft)": "en-GB-SoniaNeural",
    "Libby (Female - UK Bright/Informal Brit)": "en-GB-LibbyNeural",
    "Maisie (Female - Informal Brit)": "en-GB-MaisieNeural",
    "Anike (Female - African/NG)": "en-NG-EzinneNeural",
    "Natasha (Female - AU Aboriginal Hint)": "en-AU-NatashaNeural",
    "Fatima (Female - Arabic/UAE)": "ar-AE-FatimaNeural",
    "Yan (Female - Asian/HK)": "zh-HK-HiuGaaiNeural",
    "Clara (Female - US Latino Hint)": "es-US-PalomaNeural",
    "Ava (Female - US Gen-Z/AA Hint)": "en-US-AvaNeural",
    "Emma (Female - Casual US)": "en-US-EmmaNeural",
    "Asilia (Female - African/KE)": "en-KE-AsiliaNeural",
    "Imani (Female - East African)": "en-TZ-ImaniNeural",
    "Leah (Female - Southern African)": "en-ZA-LeahNeural",
    "Ramona (Female - Latina)": "es-DO-RamonaNeural",
    "Belkys (Female - Afro Latina/Latina)": "es-CU-BelkysNeural",
    "Svetlana (Female - Russian)": "ru-RU-SvetlanaNeural",
    "Denise (Female - French)": "fr-FR-DeniseNeural",
    "Katja (Female - German)": "de-DE-KatjaNeural",
}

VOICE_LABELS = {v: k for k, v in FREE_VOICES.items()}


# --- INITIALIZE STATE ---
def init_state():
    if "script_text" not in st.session_state:
        st.session_state.script_text = PLACEHOLDER_SCRIPT
    if "undo_stack" not in st.session_state:
        st.session_state.undo_stack = []
    if "redo_stack" not in st.session_state:
        st.session_state.redo_stack = []
    if "edit_history" not in st.session_state:
        st.session_state.edit_history = []
    if "voice_map" not in st.session_state:
        st.session_state.voice_map = {"NARRATOR": "en-GB-RyanNeural"}
    if "playing" not in st.session_state:
        st.session_state.playing = False
    if "current_line_idx" not in st.session_state:
        st.session_state.current_line_idx = 0
    if "last_active_role" not in st.session_state:
        st.session_state.last_active_role = "NARRATOR"
    if "editor_version" not in st.session_state:
        st.session_state.editor_version = 0
    if "trigger_scroll" not in st.session_state:
        st.session_state.trigger_scroll = False
    if "has_user_started_playback" not in st.session_state:
        st.session_state.has_user_started_playback = False


init_state()


# --- LOGIC: AUDIO ENGINE ---
async def generate_voice_bytes(text, voice_id, rate="+0%"):
    try:
        communicate = edge_tts.Communicate(text, voice_id, rate=rate)
        audio_data = b""
        async for chunk in communicate.stream():
            if chunk["type"] == "audio":
                audio_data += chunk["data"]
        return audio_data
    except Exception as e:
        st.error(f"Voice Error ({voice_id}): {e}")
        return None


def normalize_script_spacing(text, leading_blank_lines=0):
    text = re.sub(r"\n{3,}", "\n\n", text).strip()
    return ("\n" * leading_blank_lines) + text


def guess_gender(name):
    """
    Intelligently assigns voices based on gender hints, common names,
    and broader culturally-informed naming patterns.
    """
    males = [v for k, v in FREE_VOICES.items() if "(Male" in k]
    females = [v for k, v in FREE_VOICES.items() if "(Female" in k]

    name_up = name.upper().strip()
    name_clean = re.sub(r'[^A-Z\s-]', '', name_up).strip()
    first_token = name_clean.split()[0] if name_clean.split() else name_clean

    fem_names = {
        "SARAH", "SARAI", "CHLOE", "EMILY", "RACHEL", "MEGAN", "AMY",
        "LAUREN", "MICHELLE", "CLAIRE", "ALICE", "BETH", "IVY", "NAOMI",
        "AISHA", "AISHAH", "ALIYAH", "AALIYAH", "AMARA", "IMANI", "NIA",
        "ZURI", "AYANA", "AYANNA", "KIARA", "KIERA", "KIRA", "TAMERA",
        "TAMIKA", "MONIQUE", "DANIELLE", "JANELLE", "JASMINE", "JAZMINE",
        "JADA", "KAYLA", "KEISHA", "KESHIA", "LATASHA", "LASHAWN", "SHAWNDA",
        "SHAKIRA", "SHANICE", "TIARA", "TAMERA", "TAMARA", "MYA", "NALA",
        "ZORA", "BREANNA", "BRIANA", "DESTINY", "EBONY", "KENDRA", "NIKKI",
        "OMARIA", "SAMAYA", "SANIYA", "TANIYA", "TRINITY", "ZANIYAH"
    }

    male_names = {
        "MALIK", "JAMAL", "JAMAAL", "KAREEM", "KARIM", "AMARI", "OMAR",
        "DARIUS", "MARCUS", "MALCOLM", "ISAIAH", "ELIJAH", "JALEN", "JAYLEN",
        "JALIN", "JABARI", "JELANI", "DEVON", "DEVIN", "DEANDRE", "DONTE",
        "DARNELL", "TERRELL", "TREVON", "TREVOR", "TYRELL", "TYREE", "TYRONE",
        "LAMAR", "LAMONT", "MARQUIS", "MARCELLUS", "MICAH", "NAJEE", "NASIR",
        "RASHEED", "RAHEEM", "TAJ", "TAHEEM", "XAVIER", "ZION", "KHALIL",
        "KHARI", "KAIRO", "JOAQUIN", "DESHAWN", "DAQUAN", "DAKWAN", "QUINTON",
        "KENDRICK", "KOBE", "KOBI", "TRISTAN", "TARIQ", "TAREK", "AHMAD",
        "AHMAUD", "KAI", "JORDAN", "COREY", "CORY", "BRYSON", "JERMAINE",
        "DEMETRIUS", "ANTWON", "ANTWAN", "ANTOINE", "KEON", "KEENAN", "OMARI"
    }

    if first_token in fem_names or name_clean in fem_names:
        return random.choice(females)

    if first_token in male_names or name_clean in male_names:
        return random.choice(males)

    if any(suffix in name_clean for suffix in ["VIK", "OV", "SLAV", "IM"]):
        return "ru-RU-DmitryNeural"
    if any(suffix in name_clean for suffix in ["OVA", "INA", "YANA"]):
        return "ru-RU-SvetlanaNeural"

    if any(suffix in name_clean for suffix in ["JEAN", "LUC", "PIERRE", "ZAMORA"]):
        return "fr-FR-HenriNeural"

    if any(suffix in name_clean for suffix in ["HELM", "RICHT", "BURG"]):
        return "de-DE-ConradNeural"

    fem_hints = [
        "ARIA", "JENNY", "SONIA", "MARIAN", "GIRL", "WOMAN", "SISTER",
        "MOTHER", "QUEEN", "LADY", "STYLIST", "MIRA", "MS", "MISS"
    ]
    masc_hints = [
        "GUY", "MAN", "BOY", "BROTHER", "FATHER", "KING", "LORD",
        "HENCHMAN", "OFFICER", "GUARD", "SUIT", "DRIVER", "FIREFIGHTER",
        "MR", "SIR"
    ]

    if any(hint in name_clean for hint in fem_hints) or first_token.endswith("A"):
        return random.choice(females)

    if any(hint in name_clean for hint in masc_hints):
        return random.choice(males)

    return random.choice(males + females)


def extract_characters(script_text):
    potential_names = re.findall(
        r"^(?!INT\.|EXT\.|SCENE|ACT)[A-Z][A-Z\d\s\.]+$",
        script_text,
        re.MULTILINE
    )

    exclude = [
        "INT.", "EXT.", "CUT TO:", "FADE IN:", "FADE OUT:", "CONTINUED:",
        "V.O.", "O.C.", "THE END", "ACT ONE", "ACT TWO", "ACT THREE",
        "SCENE", "TITLE", "CARD", "PAGE", "DAY", "NIGHT", "DISSOLVE TO:",
        "MOMENTS LATER", "LATER", "PROLOGUE", "EPILOGUE", "FADE TO:",
        "FADE TO BLACK.", "BEAT.", "MATCH CUT:", "JUMP CUT:", "BACK TO:"
    ]

    clean_names = []
    for n in potential_names:
        name = n.strip()
        if (
            name not in exclude
            and not name.endswith(".")
            and len(name) > 1
            and not re.match(r"^\d+\.$", name)
        ):
            clean_names.append(name)

    return sorted(list(set(clean_names)))


def get_docx_download(text):
    doc = Document()
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Courier New"
    font.size = Pt(12)

    for line in text.split("\n"):
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.name = "Courier New"
        run.font.size = Pt(12)

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()


def build_scroll_js(target_raw_line):
    return f"""
    <script>
    setTimeout(function() {{
        var textArea = window.parent.document.querySelector('textarea');
        if (textArea) {{
            var text = textArea.value;
            var lines = text.split('\\n');
            var charPos = 0;
            for (var i = 0; i < Math.min({target_raw_line}, lines.length); i++) {{
                charPos += lines[i].length + 1;
            }}

            textArea.focus({{ preventScroll: true }});
            textArea.setSelectionRange(
                charPos,
                charPos + (lines[{target_raw_line}] ? lines[{target_raw_line}].length : 0)
            );

            var lineHeight = 40;

            // Keep the highlighted line well below the top edge
            var internalOffset = ({target_raw_line} * lineHeight) - (textArea.clientHeight * 0.42);

            textArea.scrollTo({{
                top: Math.max(0, internalOffset),
                behavior: 'smooth'
            }});
        }}
    }}, 50);
    </script>
    """


def scroll_enabled():
    return st.session_state.has_user_started_playback


# --- CUSTOM CSS FOR STYLING ---
st.markdown("""
    <style>
    [data-testid="stHeader"] {
        background-color: rgba(0,0,0,0) !important;
        color: white !important;
    }

    .compact-header {
        margin-top: 35px !important;
        z-index: 999999;
        font-family: 'Courier New', Courier, monospace;
        background-color: #262730;
        padding: 10px 20px;
        border-radius: 4px;
        color: #ffd600;
        font-size: 16px;
        margin-bottom: 10px;
        display: flex;
        justify-content: space-between;
        align-items: center;
        border: 1px solid #ffd600;
        position: relative;
    }

    .block-container {
        padding-top: 2rem !important;
        padding-bottom: 0rem !important;
        max-width: 95% !important;
    }

    footer {visibility: hidden;}
    #MainMenu {visibility: hidden;}

    .sticky-prompter {
        position: sticky;
        top: 60px;
        z-index: 999;
        background-color: #0e1117;
        padding-bottom: 15px;
        border-bottom: 1px solid #333;
    }

    div[data-testid="stTextArea"] textarea {
        font-family: 'Courier New', Courier, monospace !important;
        background-color: #ffffff !important;
        color: #000000 !important;
        font-size: 18px !important;
        line-height: 1.6 !important;
        padding: 40px !important;
    }

    .performance-monitor {
        background-color: #1e1e1e;
        color: #ffffff;
        padding: 15px;
        border-radius: 8px;
        border-left: 10px solid #ffd600;
        font-family: 'Courier New', Courier, monospace;
    }

    .active-line {
        color: #ffd600;
        font-weight: bold;
        font-size: 24px;
        display: block;
        margin-top: 5px;
    }

    .stats-badge {
        background-color: #ffd600;
        color: #000;
        padding: 2px 10px;
        border-radius: 4px;
        font-weight: bold;
    }

    @keyframes pulse-yellow {
        0% { box-shadow: 0 0 0 0 rgba(255, 214, 0, 0.7); }
        70% { box-shadow: 0 0 0 10px rgba(255, 214, 0, 0); }
        100% { box-shadow: 0 0 0 0 rgba(255, 214, 0, 0); }
    }

    .mobile-hint {
        display: none;
        background-color: #ffd600;
        color: #000;
        padding: 12px;
        border-radius: 4px;
        text-align: center;
        font-family: 'Courier New', Courier, monospace;
        font-weight: bold;
        margin-top: 10px;
        margin-bottom: 15px;
        border: 2px solid #000;
        animation: pulse-yellow 2s infinite;
    }

    @media (max-width: 768px) {
        .mobile-hint {
            display: block;
        }
    }

    .main .block-container {
        padding-bottom: 10rem !important;
    }

    div[data-testid="stTextArea"] {
        margin-bottom: 150px !important;
    }

    div.stButton > button[kind="primary"] {
        background-color: #ffd600 !important;
        color: #000000 !important;
        border: none !important;
        font-weight: bold !important;
    }

    div.stButton > button[kind="primary"]:hover {
        background-color: #ccac00 !important;
        color: #000000 !important;
    }

    div[data-testid="stSlider"] div[data-baseweb="slider"] > div:first-child > div:nth-child(2) {
        background-color: #ffd600 !important;
    }

    div[role="slider"] {
        background-color: #ffffff !important;
        border: 2px solid #ffd600 !important;
    }
    </style>
""", unsafe_allow_html=True)

word_count = len(st.session_state.script_text.split()) if st.session_state.script_text else 0

st.markdown(f'''
    <div class="compact-header">
        <span>🎭 NARRATIVE SOUNDSTAGE</span>
        <span>WORDS: <span class="stats-badge">{word_count}</span></span>
    </div>
''', unsafe_allow_html=True)

st.markdown(
    '<div class="mobile-hint">⬅️ OPEN SIDEBAR FOR STUDIO CONTROLS</div>',
    unsafe_allow_html=True
)


# --- SIDEBAR ---
with st.sidebar:
    st.title("🎬 Studio Controls")
    st.subheader("⏱️ Reading Controls")

    pause_buffer = st.slider("Line Pause Buffer (Secs)", 0.0, 2.0, 1.3, 0.1)
    speed_factor = st.slider("Playback Speed", 0.5, 2.0, 1.4, 0.1, key="speed")
    rate_str = f"{int((speed_factor - 1) * 100):+d}%"

    col1, col2 = st.columns(2)
    with col1:
        if st.session_state.playing:
            if st.button("⏸️ PAUSE", use_container_width=True):
                st.session_state.playing = False
                st.rerun()
        else:
            if st.button("▶️ PLAY", use_container_width=True, type="primary"):
                st.session_state.playing = True
                st.session_state.has_user_started_playback = True
                st.rerun()

    with col2:
        if st.button("🔄 RESTART", use_container_width=True):
            st.session_state.current_line_idx = 0
            st.session_state.playing = False
            st.rerun()

    h_col1, h_col2 = st.columns(2)
    with h_col1:
        if st.button("↩️ UNDO", use_container_width=True, disabled=len(st.session_state.undo_stack) < 1):
            st.session_state.redo_stack.append(st.session_state.script_text)
            st.session_state.script_text = st.session_state.undo_stack.pop()
            st.session_state.editor_version += 1
            st.session_state.playing = False
            st.rerun()

    with h_col2:
        if st.button("↪️ REDO", use_container_width=True, disabled=len(st.session_state.redo_stack) < 1):
            st.session_state.undo_stack.append(st.session_state.script_text)
            st.session_state.script_text = st.session_state.redo_stack.pop()
            st.session_state.editor_version += 1
            st.session_state.playing = False
            st.rerun()

    if st.session_state.script_text:
        if st.button("🖱️ EDIT CURRENT LINE", use_container_width=True):
            st.session_state.playing = False
            st.session_state.trigger_scroll = True
            st.rerun()

    st.subheader("🔍 Find & Replace")
    f_text = st.text_input("Find text...")
    r_text = st.text_input("Replace with...")

    if st.button("Apply Replace", use_container_width=True):
        if f_text:
            st.session_state.undo_stack.append(st.session_state.script_text)
            st.session_state.script_text = st.session_state.script_text.replace(f_text, r_text)
            st.session_state.editor_version += 1
            st.session_state.playing = False
            st.rerun()

    st.subheader("📍 Navigation")
    jump_line = st.number_input("Jump to Block #", min_value=1, value=st.session_state.current_line_idx + 1)
    if st.button("🚀 Jump", use_container_width=True):
        st.session_state.current_line_idx = int(jump_line) - 1
        st.session_state.playing = False
        st.rerun()

    audio_engine_slot = st.empty()

    if st.session_state.script_text:
        docx_data = get_docx_download(st.session_state.script_text)
        st.download_button("💾 Export .docx", docx_data, "Script_Updated.docx", use_container_width=True)

    st.markdown("---")
    st.markdown("Follow me on:")
    st.markdown(f"LinkedIn $\\rightarrow$ [Ida Akiwumi](https://www.linkedin.com/in/idaa11)")
    st.markdown(f"""
        **Developed by {__author__}**, *Product Architect & Narrative Strategist* Specializing in AI-assisted workflows for Film, Gaming, and Data Storytelling.
    """)


# --- MAIN INTERFACE ---
uploaded_file = st.file_uploader("Upload Script", type=["docx"], key="file_input", label_visibility="collapsed")

if uploaded_file:
    doc = Document(uploaded_file)
    raw_text = "\n".join([p.text for p in doc.paragraphs])
    normalized_text = normalize_script_spacing(
        raw_text,
        leading_blank_lines=UPLOAD_TOP_PADDING_LINES
    )

    if normalized_text != st.session_state.script_text:
        st.session_state.undo_stack.append(st.session_state.script_text)
        st.session_state.script_text = normalized_text
        st.session_state.current_line_idx = 0
        st.session_state.editor_version += 1
        st.session_state.playing = False
        st.rerun()

if st.session_state.script_text:
    cast_list = extract_characters(st.session_state.script_text)

    with st.expander("👤 Casting Office"):
        c1, c2 = st.columns(2)

        current_n_id = st.session_state.voice_map.get("NARRATOR", "en-GB-RyanNeural")
        n_label = VOICE_LABELS.get(current_n_id, "Eric (Male - UK Narrative)")
        chosen_n = st.selectbox(
            "Narrator Voice",
            list(FREE_VOICES.keys()),
            index=list(FREE_VOICES.keys()).index(n_label)
        )
        st.session_state.voice_map["NARRATOR"] = FREE_VOICES[chosen_n]

        for i, name in enumerate(cast_list):
            if name not in st.session_state.voice_map:
                st.session_state.voice_map[name] = guess_gender(name)

            with (c1 if i % 2 == 0 else c2):
                current_v_id = st.session_state.voice_map[name]
                v_label = VOICE_LABELS.get(current_v_id, "Guy (Male - US)")
                chosen_v = st.selectbox(
                    f"Role: {name}",
                    list(FREE_VOICES.keys()),
                    index=list(FREE_VOICES.keys()).index(v_label),
                    key=f"v_sel_{name}"
                )
                st.session_state.voice_map[name] = FREE_VOICES[chosen_v]

    raw_split = st.session_state.script_text.split("\n")
    lines = []
    for i, line in enumerate(raw_split):
        stripped = line.strip()
        if not stripped:
            continue

        is_dialogue = False
        is_parenthetical = stripped.startswith("(") and stripped.endswith(")")

        if i > 0 and not is_parenthetical:
            prev = raw_split[i - 1].strip()
            if prev in cast_list or (prev.startswith("(") and prev.endswith(")")):
                is_dialogue = True

        lines.append({
            "text": stripped,
            "is_dialogue": is_dialogue,
            "is_parenthetical": is_parenthetical,
            "raw_line_num": i
        })

    scroll_js = ""
    if lines:
        curr_idx = min(st.session_state.current_line_idx, len(lines) - 1)
        current_line_text = lines[curr_idx]["text"]
        target_raw_line = lines[curr_idx]["raw_line_num"]

        st.markdown(f"""
            <div class="sticky-prompter">
                <div class="performance-monitor">
                    <div style="display: flex; justify-content: space-between;">
                        <small style="color:#ffd600;">BLOCK {curr_idx + 1} OF {len(lines)}</small>
                        <small style="color:#ffd600;">{"● READING" if st.session_state.playing else "○ STANDBY"}</small>
                    </div>
                    <span class="active-line">{current_line_text}</span>
                </div>
            </div>
        """, unsafe_allow_html=True)

        if scroll_enabled():
            scroll_js = build_scroll_js(target_raw_line)

        if st.session_state.trigger_scroll and scroll_enabled() and scroll_js:
            components.html(scroll_js, height=0)
            time.sleep(HIGHLIGHT_SETTLE_TIME)
            st.session_state.trigger_scroll = False
        elif st.session_state.trigger_scroll:
            st.session_state.trigger_scroll = False

    current_editor_val = st.text_area(
        "Script",
        value=st.session_state.script_text,
        height=900,
        key=f"editor_v{st.session_state.editor_version}",
        label_visibility="collapsed"
    )

    if current_editor_val != st.session_state.script_text:
        st.session_state.script_text = current_editor_val
        st.session_state.playing = False

    # --- RUNTIME (SCROLL & SOUND) ---
    if st.session_state.playing and lines:
        if st.session_state.current_line_idx < len(lines):
            if scroll_enabled() and scroll_js:
                components.html(scroll_js, height=0)
                time.sleep(HIGHLIGHT_SETTLE_TIME)

            idx = st.session_state.current_line_idx
            line_data = lines[idx]
            line_content = line_data["text"]

            if line_content in cast_list:
                st.session_state.last_active_role = line_content
                current_role = "NARRATOR"
                read_text = line_content
            elif line_data.get("is_parenthetical"):
                current_role = "NARRATOR"
                read_text = line_content
            elif line_content.isupper() and any(x in line_content for x in ["INT.", "EXT.", "DAY", "NIGHT"]):
                st.session_state.last_active_role = "NARRATOR"
                current_role = "NARRATOR"
                read_text = line_content
            elif line_data["is_dialogue"]:
                current_role = st.session_state.last_active_role
                read_text = line_content
            else:
                st.session_state.last_active_role = "NARRATOR"
                current_role = "NARRATOR"
                read_text = line_content

            v_id = st.session_state.voice_map.get(current_role, st.session_state.voice_map["NARRATOR"])

            loop = asyncio.new_event_loop()
            try:
                asyncio.set_event_loop(loop)
                audio_data = loop.run_until_complete(generate_voice_bytes(read_text, v_id, rate=rate_str))
            finally:
                loop.close()

            if audio_data:
                b64 = base64.b64encode(audio_data).decode()
                audio_tag = f'<audio autoplay="true" id="p_{time.time()}"><source src="data:audio/mp3;base64,{b64}"></audio>'
                audio_engine_slot.markdown(audio_tag, unsafe_allow_html=True)

                w_count = len(read_text.split())
                base_wait = ((max(2.2, (w_count / 140) * 60)) / speed_factor)

                if line_content in cast_list:
                    wait_time = CAST_NAME_WAIT
                elif line_data.get("is_parenthetical"):
                    wait_time = PARENTHETICAL_WAIT
                elif line_data["is_dialogue"]:
                    wait_time = base_wait + pause_buffer
                else:
                    wait_time = base_wait + (pause_buffer * 0.7)

                wait_time = max(wait_time, MIN_LINE_DISPLAY_TIME)
                time.sleep(wait_time)

            st.session_state.current_line_idx += 1
            st.rerun()
        else:
            st.session_state.playing = False
            st.session_state.current_line_idx = 0
            st.rerun()